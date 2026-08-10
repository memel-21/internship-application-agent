"""Generate approved cover letters as DOCX and PDF files."""

import re
from dataclasses import dataclass
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER  # type: ignore[import-untyped]
from reportlab.lib.styles import getSampleStyleSheet  # type: ignore[import-untyped]
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer  # type: ignore[import-untyped]

from internship_agent.domain.vacancy import ApplicationStatus
from internship_agent.exceptions import DocumentGenerationError
from internship_agent.persistence.models import ApplicationRecord


@dataclass(frozen=True)
class GeneratedDocumentPaths:
    """Paths for generated cover letter documents."""

    docx_path: Path
    pdf_path: Path


def generate_cover_letter_documents(
    application: ApplicationRecord,
    *,
    output_dir: Path,
    overwrite: bool = False,
) -> GeneratedDocumentPaths:
    """Generate DOCX and PDF cover letters for an approved application."""

    if application.status != ApplicationStatus.APPROVED.value:
        raise DocumentGenerationError("Only approved applications can generate documents.")
    cover_letter_text = application.cover_letter_text
    if not cover_letter_text:
        raise DocumentGenerationError("Approved application has no cover letter text to export.")

    target_dir = (output_dir / "cover_letters").resolve()
    target_dir.mkdir(parents=True, exist_ok=True)
    filename_stem = _cover_letter_stem(application)
    docx_path = (target_dir / f"{filename_stem}.docx").resolve()
    pdf_path = (target_dir / f"{filename_stem}.pdf").resolve()
    _ensure_within_directory(target_dir, docx_path)
    _ensure_within_directory(target_dir, pdf_path)
    _ensure_can_write(docx_path, overwrite=overwrite)
    _ensure_can_write(pdf_path, overwrite=overwrite)

    _write_docx(application, docx_path, cover_letter_text)
    _write_pdf(application, pdf_path, cover_letter_text)
    return GeneratedDocumentPaths(docx_path=docx_path, pdf_path=pdf_path)


def _cover_letter_stem(application: ApplicationRecord) -> str:
    raw = f"{application.company_name}-{application.role_title}-{application.id}"
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", raw).strip("-").casefold()
    if not slug:
        raise DocumentGenerationError("Could not create a safe output filename.")
    return slug[:120]


def _ensure_within_directory(parent: Path, child: Path) -> None:
    if parent not in child.parents:
        raise DocumentGenerationError("Generated document path escaped the output directory.")


def _ensure_can_write(path: Path, *, overwrite: bool) -> None:
    if path.exists() and not overwrite:
        raise DocumentGenerationError(f"Refusing to overwrite existing document: {path}")


def _write_docx(application: ApplicationRecord, path: Path, cover_letter_text: str) -> None:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_run = heading.add_run("Cover Letter")
    heading_run.bold = True
    heading_run.font.name = "Arial"
    heading_run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"{application.company_name} - {application.role_title}")

    document.add_paragraph("")
    for block in cover_letter_text.split("\n"):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.add_run(block)

    document.save(str(path))


def _write_pdf(application: ApplicationRecord, path: Path, cover_letter_text: str) -> None:
    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72,
        title=f"Cover Letter - {application.company_name}",
    )
    styles = getSampleStyleSheet()
    story: list[Paragraph | Spacer] = [
        Paragraph("Cover Letter", styles["Title"]),
        Paragraph(f"{application.company_name} - {application.role_title}", styles["Heading2"]),
        Spacer(1, 18),
    ]
    for block in cover_letter_text.split("\n"):
        text = escape(block.strip()) or "&nbsp;"
        story.append(Paragraph(text, styles["BodyText"]))
        story.append(Spacer(1, 8))
    doc.build(story)
